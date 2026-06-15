"""
Contract generator based on real .docx templates.
Replaces variable data in template runs by searching for placeholder content.
Works for both Mahsutov and Bilim templates.
"""
import re
from io import BytesIO
from docx import Document


def number_to_words(n):
    try:
        n = int(n)
    except:
        return str(n)
    if n == 0:
        return "Ноль"
    ones = ['','один','два','три','четыре','пять','шесть','семь','восемь','девять',
            'десять','одиннадцать','двенадцать','тринадцать','четырнадцать','пятнадцать',
            'шестнадцать','семнадцать','восемнадцать','девятнадцать']
    tens = ['','','двадцать','тридцать','сорок','пятьдесят',
            'шестьдесят','семьдесят','восемьдесят','девяносто']
    hundreds = ['','сто','двести','триста','четыреста','пятьсот',
                'шестьсот','семьсот','восемьсот','девятьсот']

    def chunk(num, fem=False):
        if num == 0: return ''
        parts = []
        h = num // 100
        r = num % 100
        if h: parts.append(hundreds[h])
        if r < 20:
            if r:
                w = ones[r]
                if fem and r == 1: w = 'одна'
                if fem and r == 2: w = 'две'
                parts.append(w)
        else:
            parts.append(tens[r // 10])
            o = r % 10
            if o:
                w = ones[o]
                if fem and o == 1: w = 'одна'
                if fem and o == 2: w = 'две'
                parts.append(w)
        return ' '.join(parts)

    parts = []
    millions = n // 1_000_000
    thousands = (n % 1_000_000) // 1_000
    remainder = n % 1_000
    if millions:
        w = chunk(millions)
        m = millions % 100
        s = 'миллионов' if (11<=m<=19 or m%10 in(0,5,6,7,8,9)) else ('миллион' if m%10==1 else 'миллиона')
        parts.append(f"{w} {s}")
    if thousands:
        w = chunk(thousands, True)
        m = thousands % 100
        s = 'тысяч' if (11<=m<=19 or m%10 in(0,5,6,7,8,9)) else ('тысяча' if m%10==1 else 'тысячи')
        parts.append(f"{w} {s}")
    if remainder:
        parts.append(chunk(remainder))
    r = ' '.join(parts).strip()
    return r[0].upper() + r[1:]


def fmt_num(val):
    try:
        n = int(str(val).replace(' ','').replace(',',''))
        return f"{n:,}".replace(",", " "), number_to_words(n), n
    except:
        return str(val), '', 0


def parse_date(s):
    months = ['января','февраля','марта','апреля','мая','июня',
              'июля','августа','сентября','октября','ноября','декабря']
    m = re.match(r'(\d{1,2})[./](\d{1,2})[./](\d{4})', str(s))
    if m:
        d, mo, y = m.groups()
        return {
            'day': d.zfill(2),
            'month_num': mo.zfill(2),
            'month_name': months[int(mo)-1],
            'year': y,
            'year_last': y[3],
        }
    return None


def is_mahsutov(paras):
    """Detect template type by run count in P4"""
    return len(paras[4].runs) > 20


def generate_contract(data, template_path):
    doc = Document(template_path)

    num = data.get('contract_num', '')
    cdate = data.get('contract_date', '')
    course = data.get('course', '')
    df = data.get('date_from', '')
    dt = data.get('date_to', '')
    branch = data.get('branch', '')
    total = data.get('total_amount', '')
    month_amt = data.get('month_amount', '')
    disc = data.get('discount_amount', '')
    pfio = data.get('parent_fio', '')
    piin = data.get('parent_iin', '')
    pdoc = data.get('parent_doc_num', '')
    pdoc_date = data.get('parent_doc_date', '')
    pphone = data.get('parent_phone', '')
    cfio = data.get('child_fio', '')
    ciin = data.get('child_iin', '')

    cd = parse_date(cdate)
    dfd = parse_date(df)
    dtd = parse_date(dt)
    pdd = parse_date(pdoc_date)

    total_str, total_words, _ = fmt_num(total)
    month_str, month_words, _ = fmt_num(month_amt)
    disc_str, disc_words, _ = fmt_num(disc)

    paras = doc.paragraphs
    mahsutov = is_mahsutov(paras)

    # === P1: Contract number ===
    p1 = paras[1]
    if len(p1.runs) > 1:
        p1.runs[1].text = num

    # === P2: Date line ===
    # Both templates: run7='«', run8=blank(day), run9='»', run11=month, run12=year+г
    if cd:
        p2 = paras[2]
        r = p2.runs
        # run7 contains «, run8=day placeholder, run9=», run11=month, run12=year
        if len(r) > 12:
            r[8].text = cd['day']
            r[11].text = cd['month_num']
            if mahsutov:
                r[12].text = ' 202'
                if len(r) > 13: r[13].text = cd['year_last']
            else:
                r[12].text = ' ' + cd['year'] + ' г. '

    # === P4: Parent data — search by content ===
    p4 = paras[4]
    r4 = p4.runs
    for ri, run in enumerate(r4):
        t = run.text
        # FIO placeholder — underscores after citizen text
        if '____' in t and ri > 8 and 'далее' not in t and 'адресу' not in t:
            run.text = pfio + ','
            break

    for ri, run in enumerate(r4):
        if '\nИИН' in run.text:
            if ri + 1 < len(r4):
                r4[ri+1].text = piin + ', '
            break

    for ri, run in enumerate(r4):
        if 'удостоверение личности №' in run.text:
            run.text = 'удостоверение личности № ' + pdoc
            # Clear extra underscore run
            if ri+1 < len(r4) and '_' in r4[ri+1].text:
                r4[ri+1].text = ''
            break

    for ri, run in enumerate(r4):
        if ', выдано «' in run.text:
            run.text = ', выдано «'
            if pdd:
                if ri+1 < len(r4): r4[ri+1].text = pdd['day'] + '»'
                if ri+2 < len(r4): r4[ri+2].text = ' ' + pdd['month_num'] + ' 20' + pdd['year'][2:]
            else:
                if ri+1 < len(r4): r4[ri+1].text = '»'
                if ri+2 < len(r4): r4[ri+2].text = ' ____'
            break

    # === P5: Child data ===
    p5 = paras[5]
    r5 = p5.runs
    for ri, run in enumerate(r5):
        if '____' in run.text and 'Слушатель' not in run.text:
            run.text = ': ' + cfio + ','
            break
    for ri, run in enumerate(r5):
        if run.text == 'ИИН ':
            if ri+1 < len(r5):
                r5[ri+1].text = ciin + ', '
            break

    # === P7: Course ===
    p7 = paras[7]
    if len(p7.runs) > 2:
        p7.runs[2].text = course

    # === P8: Period ===
    # Both: run0=text(no day), run1='» ', run2=month+year, run3=г, run4=по«, run5=day_to, run6='» ', run7=month_to, run8=year_to
    if dfd and dtd:
        p8 = paras[8]
        r8 = p8.runs
        # Insert from-day into run0 text
        r8[0].text = f'1.2. Период оказания Услуг по настоящему Договору с «{dfd["day"]}» '
        r8[1].text = ''
        if mahsutov:
            r8[2].text = dfd['month_name'] + ' 202'
            r8[3].text = dfd['year_last']
            r8[4].text = ' '
            r8[5].text = 'г. '
            r8[6].text = 'по «'
            r8[7].text = dtd['day']
            r8[8].text = '» '
            r8[9].text = dtd['month_name'] + ' '
            r8[10].text = dtd['year'] + ' г.'
            if len(r8) > 11:
                r8[11].text = ', если иное не согласовано Сторонами в письменной форме.'
        else:
            r8[2].text = dfd['month_name'] + ' ' + dfd['year'] + ' '
            r8[3].text = 'г. '
            r8[4].text = 'по «'
            r8[5].text = dtd['day']
            r8[6].text = '» '
            r8[7].text = dtd['month_name'] + ' '
            r8[8].text = dtd['year'] + ' г.'
            if len(r8) > 9:
                r8[9].text = ', если иное не согласовано Сторонами в письменной форме.'

    # === P16: Branch 1.6 ===
    p16 = paras[16]
    if len(p16.runs) > 2:
        p16.runs[2].text = ', ' + branch + '.'

    # === TABLE 0: Requisites ===
    tbl0 = doc.tables[0]
    cell_parent = tbl0.rows[0].cells[1]
    cp = cell_parent.paragraphs

    def set_cell_para(paras_list, idx, text):
        if idx < len(paras_list):
            p = paras_list[idx]
            if p.runs:
                p.runs[0].text = text
                for r in p.runs[1:]: r.text = ''

    set_cell_para(cp, 1, 'ФИО: ' + pfio)
    set_cell_para(cp, 2, 'ИИН  ' + piin)
    set_cell_para(cp, 3, 'удостоверение личности № ' + pdoc)
    if pdd:
        set_cell_para(cp, 4, f'выдано МВД РК от «{pdd["day"]}» {pdd["month_num"]} 20{pdd["year"][2:]}г.')
    if pphone and len(cp) > 6:
        set_cell_para(cp, 6, 'Конт.тел.: ' + pphone)

    # === TABLE 1: Appendix 1 ===
    tbl1 = doc.tables[1]

    # Row 0: Total amount + dates
    cell_total = tbl1.rows[0].cells[2]
    cp0 = cell_total.paragraphs
    if dfd and dtd and len(cp0) > 0:
        r = cp0[0].runs
        if mahsutov:
            # run2=day_from, run3=month_from+202, run4=year_last, run8=day_to, run9=month_to+year
            if len(r) > 9:
                r[2].text = dfd['day'] + '» '
                r[3].text = dfd['month_name'] + ' 202'
                r[4].text = dfd['year_last']
                r[8].text = dtd['day'] + '» '
                r[9].text = dtd['month_name'] + ' ' + dtd['year'] + ' '
        else:
            # run2=day_from, run3=month_from+year, run6=day_to, run7=month_to+year
            if len(r) > 7:
                r[2].text = dfd['day'] + '» '
                r[3].text = dfd['month_name'] + ' ' + dfd['year'] + ' '
                r[6].text = dtd['day'] + '» '
                r[7].text = dtd['month_name'] + ' ' + dtd['year'] + ' '

    if total and len(cp0) > 1:
        p1r = cp0[1].runs
        if p1r:
            p1r[0].text = f'{total_str} ({total_words})'
            for r in p1r[1:]: r.text = ''

    # Row 1: Month amount
    # run2 = '124 000 (Сто двадцать четыре тысячи) тенге.'
    if month_amt:
        cell_month = tbl1.rows[1].cells[2]
        mp = cell_month.paragraphs
        if mp:
            mr = mp[0].runs
            if len(mr) > 2:
                mr[2].text = f'{month_str} ({month_words}) тенге.'
                for r in mr[3:]: r.text = ''

    # Row 2: Discount
    if disc:
        cell_disc = tbl1.rows[2].cells[2]
        dp = cell_disc.paragraphs
        # P1: total reference, P2: discount amount
        if len(dp) > 2:
            dr1 = dp[1].runs
            if len(dr1) > 2:
                dr1[2].text = f'{total_str} тенге'
                for r in dr1[3:]: r.text = ''
            dr2 = dp[2].runs
            if dr2:
                dr2[0].text = f'{disc_str} ({disc_words}) тенге.'
                for r in dr2[1:]: r.text = ''

    # Row 4: Branch
    cell_branch = tbl1.rows[4].cells[2]
    br = cell_branch.paragraphs[0].runs if cell_branch.paragraphs else []
    if len(br) > 2:
        br[2].text = ', ' + branch

    # === P153: Appendix 1 date ===
    if cd:
        p153 = paras[153]
        r153 = p153.runs
        if len(r153) > 1:
            r153[0].text = f'от «{cd["day"]}» {cd["month_num"]} 202'
            r153[1].text = cd['year_last']
            if len(r153) > 2: r153[2].text = ' года'
        elif len(r153) == 1:
            r153[0].text = f'от «{cd["day"]}» {cd["month_num"]} {cd["year"]} года'

        # === P157: App1 city+date ===
        p157 = paras[157]
        r157 = p157.runs
        if mahsutov:
            # run3=year_last, run4=' года' — need to add date
            if len(r157) > 2:
                r157[2].text = f'   «{cd["day"]}» {cd["month_num"]} 202'
                if len(r157) > 3: r157[3].text = cd['year_last']
                if len(r157) > 4: r157[4].text = ' года'
        else:
            # run2 = spaces, append date
            if len(r157) > 2:
                r157[2].text = f'   «{cd["day"]}» {cd["month_num"]} {cd["year"]} года'

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()

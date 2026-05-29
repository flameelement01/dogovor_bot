import os
import re
import logging
from io import BytesIO
from datetime import date

from ocr import extract_from_pdfs
from contract_generator import generate_contract as generate_from_template

import fitz
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup, ReplyKeyboardRemove
from telegram.ext import (
    Application, CommandHandler, MessageHandler, CallbackQueryHandler,
    ConversationHandler, filters, ContextTypes
)

logging.basicConfig(format='%(asctime)s - %(levelname)s - %(message)s', level=logging.INFO)
logger = logging.getLogger(__name__)

BOT_TOKEN = os.getenv("BOT_TOKEN", "YOUR_BOT_TOKEN")

IP_DATA = {
    "mahsutov": {
        "label": "ИП «Махсутов»", "full_name": "Махсутов",
        "director": "Махсутов Адиль Шынаралович",
        "notice": "KZ62UWQ04216925", "notice_date": "29 ноября 2022",
        "iin": "950713301438", "address": "г. Шымкент, пр. Кунаева 95/1",
        "bank": "АО «Kaspi Bank»", "account": "KZ07722S000020485146",
        "phone": "+7 708 287 0264",
    },
    "bilim": {
        "label": "ИП «Білім Орталығы»", "full_name": "Білім Орталығы",
        "director": "Қалдыораз Абылайхан Аққалиұлы",
        "notice": "KZ24UWQ05575415", "notice_date": "16 октября 2023",
        "iin": "960201300071", "address": "г. Шымкент, ул. Желтоксана 35",
        "bank": "АО «Kaspi Bank»", "account": "KZ97722S000030494217",
        "phone": "+7 708 287 0264",
    }
}

(
    SELECT_IP, UPLOAD_PARENT_PDF, UPLOAD_CHILD_PDF,
    ENTER_CONTRACT_NUM, SELECT_COURSE, ENTER_COURSE_CUSTOM,
    ENTER_DATE_FROM, ENTER_DATE_TO, SELECT_BRANCH,
    ENTER_TOTAL_AMOUNT, ENTER_MONTH_AMOUNT, ENTER_DISCOUNT_AMOUNT,
    CONFIRM,
) = range(13)

COURSES = {
    "course_top": "ТОП школы (НИШ, БИЛ, РФМШ)",
    "course_ent": "ЕНТ",
    "course_level": "Повышение уровня знаний",
    "course_ind": "Индивидуальные уроки",
}

BRANCHES = {
    "branch_kunaeva": "пр. Кунаева 95/1",
    "branch_zheltoksana": "ул. Желтоксана 35",
    "branch_shayakhmetova": "ул. Шаяхметова 39",
}

# ==================== HELPERS ====================

def number_to_words(n):
    try:
        n = int(n)
    except:
        return str(n)
    if n == 0:
        return "Ноль"
    ones = ['','один','два','три','четыре','пять','шесть','семь','восемь','девять',
            'десять','одиннадцать','двенадцать','тринадцать','четырнадцать','пятнадцать',
            'шестнадцать','семнадцать','восемнадцать','девятнадцать']
    tens = ['','','двадцать','тридцать','сорок','пятьдесят',
            'шестьдесят','семьдесят','восемьдесят','девяносто']
    hundreds = ['','сто','двести','триста','четыреста','пятьсот',
                'шестьсот','семьсот','восемьсот','девятьсот']

    def chunk(num, fem=False):
        if num == 0: return ''
        parts = []
        h = num // 100
        r = num % 100
        if h: parts.append(hundreds[h])
        if r < 20:
            if r:
                w = ones[r]
                if fem and r == 1: w = 'одна'
                if fem and r == 2: w = 'две'
                parts.append(w)
        else:
            parts.append(tens[r // 10])
            o = r % 10
            if o:
                w = ones[o]
                if fem and o == 1: w = 'одна'
                if fem and o == 2: w = 'две'
                parts.append(w)
        return ' '.join(parts)

    parts = []
    millions = n // 1_000_000
    thousands = (n % 1_000_000) // 1_000
    remainder = n % 1_000
    if millions:
        w = chunk(millions)
        m = millions % 100
        s = 'миллионов' if (11<=m<=19 or m%10 in(0,5,6,7,8,9)) else ('миллион' if m%10==1 else 'миллиона')
        parts.append(f"{w} {s}")
    if thousands:
        w = chunk(thousands, True)
        m = thousands % 100
        s = 'тысяч' if (11<=m<=19 or m%10 in(0,5,6,7,8,9)) else ('тысяча' if m%10==1 else 'тысячи')
        parts.append(f"{w} {s}")
    if remainder:
        parts.append(chunk(remainder))
    r = ' '.join(parts).strip()
    return r[0].upper() + r[1:]


def fmt_amount(val):
    try:
        n = int(str(val).replace(' ','').replace(',',''))
        return f"{n:,}".replace(",", " ") + f" ({number_to_words(n)}) тенге"
    except:
        return str(val)


def parse_date(s):
    months = ['января','февраля','марта','апреля','мая','июня',
              'июля','августа','сентября','октября','ноября','декабря']
    m = re.match(r'(\d{1,2})[./](\d{1,2})[./](\d{4})', str(s))
    if m:
        d, mo, y = m.groups()
        return int(d), months[int(mo)-1], y
    return None, s, ''


def fmt_period(s):
    d, mn, y = parse_date(s)
    return f"«{d:02d}» {mn} {y}" if d else s


def fmt_date_str(s):
    d, mn, y = parse_date(s)
    return f"«{d}» {mn} {y} г." if d else s


def fmt_date_short(s):
    m = re.match(r'(\d{1,2})[./](\d{1,2})[./](\d{4})', str(s))
    if m:
        d, mo, y = m.groups()
        return f"{int(d):02d}.{int(mo):02d}.{y}"
    return s


# ==================== DOCX ====================

def set_font(run, size=12, bold=False, underline=False, italic=False):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.underline = underline
    run.font.italic = italic
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rFonts.set(qn('w:cs'), 'Times New Roman')
    rPr.insert(0, rFonts)


def para(doc, parts, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=True, space_after=6, indent_cm=1.25):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = Pt(14)
    if indent:
        p.paragraph_format.first_line_indent = Cm(indent_cm)
    if isinstance(parts, str):
        parts = [(parts, False, False)]
    for item in parts:
        if len(item) == 2:
            text, bold = item
            ul = False
        else:
            text, bold, ul = item
        r = p.add_run(text)
        set_font(r, bold=bold, underline=ul)
    return p


def add_tbl_cell(cell, lines):
    cell.paragraphs[0].clear()
    for i, (txt, bold) in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(txt)
        set_font(r, bold=bold)



# ==================== HANDLERS ====================

async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    context.user_data.clear()
    keyboard = [
        [InlineKeyboardButton("ИП «Махсутов»", callback_data="ip_mahsutov")],
        [InlineKeyboardButton("ИП «Білім Орталығы»", callback_data="ip_bilim")],
    ]
    await update.message.reply_text(
        "👋 *Генератор договоров AIPLUS*\n\nВыберите ИП для договора:",
        reply_markup=InlineKeyboardMarkup(keyboard),
        parse_mode='Markdown'
    )
    return SELECT_IP


async def select_ip(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    ip_key = query.data.replace("ip_", "")
    context.user_data['ip'] = ip_key
    ip = IP_DATA[ip_key]
    await query.edit_message_text(
        f"✅ Выбрано: *{ip['label']}*\n\n📄 Отправьте PDF или фото *удостоверения личности родителя*",
        parse_mode='Markdown'
    )
    return UPLOAD_PARENT_PDF


async def upload_parent_pdf(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if update.message.photo:
        file = await update.message.photo[-1].get_file()
        context.user_data['parent_pdf'] = bytes(await file.download_as_bytearray())
    elif update.message.document:
        file = await update.message.document.get_file()
        context.user_data['parent_pdf'] = bytes(await file.download_as_bytearray())
    else:
        await update.message.reply_text("⚠️ Отправьте PDF или фото удостоверения личности родителя.")
        return UPLOAD_PARENT_PDF
    await update.message.reply_text(
        "✅ Получено!\n\n📄 Теперь отправьте PDF или фото *документа ребёнка* (свидетельство о рождении или удостоверение).",
        parse_mode='Markdown'
    )
    return UPLOAD_CHILD_PDF


async def upload_child_pdf(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if update.message.photo:
        file = await update.message.photo[-1].get_file()
        context.user_data['child_pdf'] = bytes(await file.download_as_bytearray())
    elif update.message.document:
        file = await update.message.document.get_file()
        context.user_data['child_pdf'] = bytes(await file.download_as_bytearray())
    else:
        await update.message.reply_text("⚠️ Отправьте PDF или фото документа ребёнка.")
        return UPLOAD_CHILD_PDF

    await update.message.reply_text("⏳ Читаю документы...")

    extracted = extract_from_pdfs(
        context.user_data['parent_pdf'],
        context.user_data['child_pdf'],
        os.getenv("GOOGLE_VISION_KEY", "")
    )
    context.user_data.update(extracted)

    found = [k for k in ['parent_fio','parent_iin','parent_doc_num','child_fio','child_iin'] if extracted.get(k)]
    if found:
        msg = (
            f"✅ *Извлечено из документов:*\n\n"
            f"👤 Родитель: {extracted.get('parent_fio', '—')}\n"
            f"🪪 ИИН: {extracted.get('parent_iin', '—')}\n"
            f"📋 УД №: {extracted.get('parent_doc_num', '—')}\n"
            f"📅 Дата выдачи: {extracted.get('parent_doc_date', '—')}\n\n"
            f"👧 Ребёнок: {extracted.get('child_fio', '—')}\n"
            f"🪪 ИИН: {extracted.get('child_iin', '—')}\n\n"
            f"⚠️ Проверьте данные — при необходимости исправите в конце."
        )
    else:
        msg = "⚠️ Не удалось извлечь данные автоматически — заполним вручную."

    await update.message.reply_text(msg, parse_mode='Markdown')
    await update.message.reply_text("📝 Введите *номер договора* (из AmoCRM):", parse_mode='Markdown')
    return ENTER_CONTRACT_NUM


async def enter_contract_num(update: Update, context: ContextTypes.DEFAULT_TYPE):
    context.user_data['contract_num'] = update.message.text.strip()
    today = date.today().strftime('%d.%m.%Y')
    context.user_data['contract_date'] = today

    keyboard = [
        [InlineKeyboardButton("ТОП школы (НИШ, БИЛ, РФМШ)", callback_data="course_top")],
        [InlineKeyboardButton("ЕНТ", callback_data="course_ent")],
        [InlineKeyboardButton("Повышение уровня знаний", callback_data="course_level")],
        [InlineKeyboardButton("Индивидуальные уроки", callback_data="course_ind")],
        [InlineKeyboardButton("✏️ Ввести вручную", callback_data="course_custom")],
    ]
    await update.message.reply_text(
        f"✅ Дата договора: *{today}* (сегодня)\n\n📚 Выберите *курс обучения*:",
        reply_markup=InlineKeyboardMarkup(keyboard),
        parse_mode='Markdown'
    )
    return SELECT_COURSE


async def select_course(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    if query.data == 'course_custom':
        await query.edit_message_text("✏️ Введите название курса вручную:")
        return ENTER_COURSE_CUSTOM
    context.user_data['course'] = COURSES[query.data]
    await query.edit_message_text(
        f"✅ Курс: *{COURSES[query.data]}*\n\n📅 Введите *дату начала* обучения (ДД.ММ.ГГГГ):",
        parse_mode='Markdown'
    )
    return ENTER_DATE_FROM


async def enter_course_custom(update: Update, context: ContextTypes.DEFAULT_TYPE):
    context.user_data['course'] = update.message.text.strip()
    await update.message.reply_text("📅 Введите *дату начала* обучения (ДД.ММ.ГГГГ):", parse_mode='Markdown')
    return ENTER_DATE_FROM


async def enter_date_from(update: Update, context: ContextTypes.DEFAULT_TYPE):
    context.user_data['date_from'] = update.message.text.strip()
    await update.message.reply_text("📅 Введите *дату окончания* обучения (ДД.ММ.ГГГГ):", parse_mode='Markdown')
    return ENTER_DATE_TO


async def enter_date_to(update: Update, context: ContextTypes.DEFAULT_TYPE):
    context.user_data['date_to'] = update.message.text.strip()
    keyboard = [
        [InlineKeyboardButton("пр. Кунаева 95/1", callback_data="branch_kunaeva")],
        [InlineKeyboardButton("ул. Желтоксана 35", callback_data="branch_zheltoksana")],
        [InlineKeyboardButton("ул. Шаяхметова 39", callback_data="branch_shayakhmetova")],
    ]
    await update.message.reply_text(
        "🏢 Выберите *филиал*:",
        reply_markup=InlineKeyboardMarkup(keyboard),
        parse_mode='Markdown'
    )
    return SELECT_BRANCH


async def select_branch(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    context.user_data['branch'] = BRANCHES[query.data]
    await query.edit_message_text(
        f"✅ Филиал: *{BRANCHES[query.data]}*\n\n💰 Введите *общую сумму* договора (только цифры):",
        parse_mode='Markdown'
    )
    return ENTER_TOTAL_AMOUNT


async def enter_total_amount(update: Update, context: ContextTypes.DEFAULT_TYPE):
    val = update.message.text.strip().replace(' ', '').replace(',', '')
    context.user_data['total_amount'] = val
    await update.message.reply_text(
        f"✅ {fmt_amount(val)}\n\n💰 Введите *сумму за месяц* (только цифры):",
        parse_mode='Markdown'
    )
    return ENTER_MONTH_AMOUNT


async def enter_month_amount(update: Update, context: ContextTypes.DEFAULT_TYPE):
    val = update.message.text.strip().replace(' ', '').replace(',', '')
    context.user_data['month_amount'] = val
    await update.message.reply_text(
        f"✅ {fmt_amount(val)}\n\n🎁 Введите *сумму со скидкой* (цифры или *нет* если без акции):",
        parse_mode='Markdown'
    )
    return ENTER_DISCOUNT_AMOUNT


async def enter_discount_amount(update: Update, context: ContextTypes.DEFAULT_TYPE):
    val = update.message.text.strip()
    context.user_data['discount_amount'] = '' if val.lower() in ('нет','no','-','0') else val.replace(' ','').replace(',','')
    context.user_data['schedule'] = []
    return await show_confirm(update.message, context)





async def show_confirm(message, context):
    d = context.user_data
    ip = IP_DATA[d.get('ip', 'mahsutov')]
    disc = d.get('discount_amount', '')
    disc_line = f"\n🎁 Скидка: {fmt_amount(disc)}" if disc else ""
    summary = (
        f"📋 *Проверьте данные:*\n\n"
        f"🏢 ИП: {ip['label']}\n"
        f"📄 №: {d.get('contract_num','—')}\n"
        f"📅 Дата: {d.get('contract_date','—')}\n"
        f"📚 Курс: {d.get('course','—')}\n"
        f"🗓 Период: {d.get('date_from','—')} — {d.get('date_to','—')}\n"
        f"🏢 Филиал: {d.get('branch','—')}\n\n"
        f"👤 Родитель: {d.get('parent_fio','—')}\n"
        f"🪪 ИИН: {d.get('parent_iin','—')}\n"
        f"📋 УД №: {d.get('parent_doc_num','—')}\n"
        f"📅 Дата выдачи: {d.get('parent_doc_date','—')}\n\n"
        f"👧 Ребёнок: {d.get('child_fio','—')}\n"
        f"🪪 ИИН: {d.get('child_iin','—')}\n\n"
        f"💰 Сумма: {fmt_amount(d.get('total_amount','0'))}\n"
        f"📆 В месяц: {fmt_amount(d.get('month_amount','0'))}"
        f"{disc_line}"
    )

    keyboard = [
        [InlineKeyboardButton("✅ Генерировать договор", callback_data="confirm_yes")],
        [InlineKeyboardButton("🔄 Начать заново", callback_data="confirm_no")],
    ]
    await message.reply_text(summary, reply_markup=InlineKeyboardMarkup(keyboard), parse_mode='Markdown')
    return CONFIRM


async def confirm(update: Update, context: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    await query.answer()
    if query.data == "confirm_no":
        await query.edit_message_text("🔄 Начинаем заново. Введите /start")
        return ConversationHandler.END

    await query.edit_message_text("⏳ Генерирую договор...")
    try:
        ip_key = context.user_data.get('ip', 'mahsutov')
        template = os.path.join(os.path.dirname(__file__), f'template_{ip_key}.docx')
        docx_bytes = generate_from_template(context.user_data, template)
        d = context.user_data
        filename = f"Договор_{d.get('contract_num','бн')}_{d.get('parent_fio','клиент').split()[0]}.docx"
        await query.message.reply_document(
            document=BytesIO(docx_bytes),
            filename=filename,
            caption=f"✅ *Договор готов!* `{filename}`\n\nДля нового договора — /start",
            parse_mode='Markdown'
        )
    except Exception as e:
        logger.error(f"Generation error: {e}")
        await query.message.reply_text(f"❌ Ошибка: {e}\n\nПопробуйте /start")
    return ConversationHandler.END


async def cancel(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text("❌ Отменено. Для нового договора — /start", reply_markup=ReplyKeyboardRemove())
    return ConversationHandler.END


def main():
    app = Application.builder().token(BOT_TOKEN).build()
    conv = ConversationHandler(
        entry_points=[CommandHandler("start", start)],
        states={
            SELECT_IP: [CallbackQueryHandler(select_ip, pattern="^ip_")],
            UPLOAD_PARENT_PDF: [
                MessageHandler(filters.Document.PDF | filters.Document.IMAGE | filters.PHOTO, upload_parent_pdf)
            ],
            UPLOAD_CHILD_PDF: [
                MessageHandler(filters.Document.PDF | filters.Document.IMAGE | filters.PHOTO, upload_child_pdf)
            ],
            ENTER_CONTRACT_NUM: [MessageHandler(filters.TEXT & ~filters.COMMAND, enter_contract_num)],
            SELECT_COURSE: [CallbackQueryHandler(select_course, pattern="^course_")],
            ENTER_COURSE_CUSTOM: [MessageHandler(filters.TEXT & ~filters.COMMAND, enter_course_custom)],
            ENTER_DATE_FROM: [MessageHandler(filters.TEXT & ~filters.COMMAND, enter_date_from)],
            ENTER_DATE_TO: [MessageHandler(filters.TEXT & ~filters.COMMAND, enter_date_to)],
            SELECT_BRANCH: [CallbackQueryHandler(select_branch, pattern="^branch_")],
            ENTER_TOTAL_AMOUNT: [MessageHandler(filters.TEXT & ~filters.COMMAND, enter_total_amount)],
            ENTER_MONTH_AMOUNT: [MessageHandler(filters.TEXT & ~filters.COMMAND, enter_month_amount)],
            ENTER_DISCOUNT_AMOUNT: [MessageHandler(filters.TEXT & ~filters.COMMAND, enter_discount_amount)],
            CONFIRM: [CallbackQueryHandler(confirm, pattern="^confirm_")],
        },
        fallbacks=[CommandHandler("cancel", cancel)],
    )
    app.add_handler(conv)
    print("🤖 Бот запущен!")
    app.run_polling(allowed_updates=Update.ALL_TYPES)


if __name__ == "__main__":
    main()
